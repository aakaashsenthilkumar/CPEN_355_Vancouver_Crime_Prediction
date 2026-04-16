from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

def make_table(headers, rows, best_row=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
    shade_row(tbl.rows[0], "1F497D")
    for cell in tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row_data in enumerate(rows):
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if best_row is not None and i == best_row:
                        run.font.bold = True
        if i % 2 == 0:
            shade_row(row, "EBF0FA")
    return tbl

# ── Title Page ────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CPEN 355 Final Project Report")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Predicting Safety Scores for Vancouver Neighbourhoods\nUsing Historical Crime Data")
r2.font.size = Pt(13)
r2.font.italic = True

doc.add_paragraph()
for name, sid in [("Aakaash Senthilkumar", "83091546"),
                  ("Vaibhav Ambastha",     "19919539"),
                  ("Arshvir Singh",        "62273818")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{name}  |  Student ID: {sid}")
    r.font.size = Pt(11)

doc.add_paragraph()
cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.add_run("University of British Columbia  |  CPEN 355  |  2024").font.size = Pt(10)

doc.add_page_break()

# ── Abstract ──────────────────────────────────────────────────────────────────
set_heading("Abstract")
body(
    "We propose a machine learning pipeline to predict a safety score (0–100) for Vancouver "
    "neighbourhoods using historical crime data from the Vancouver Police Department (VPD). "
    "A weighted scoring system is designed to reflect crime severity across 13 crime types, "
    "and three regression models are trained and evaluated: Linear Regression, Decision Tree, "
    "and Random Forest. Using a time-based train/test split (2003–2019 train, 2020–2023 test) "
    "to prevent data leakage, our best model — Random Forest — achieves an RMSE of 1.0021 and "
    "R² of 0.9944, outperforming Linear Regression (RMSE 1.1748) and Decision Tree (RMSE 1.0313). "
    "The trained model is deployed through a CLI query tool, a terminal chatbot, and a full GUI "
    "chatbot powered by a locally running LLM (Ollama / llama3.2), enabling natural language "
    "queries over both historical data and future year predictions."
)

doc.add_page_break()

# ── 1. Introduction ───────────────────────────────────────────────────────────
set_heading("1. Introduction")
set_heading("1.1  Motivation", level=2)
body(
    "Urban safety is a critical factor in city planning, real-estate valuation, and public "
    "awareness. Vancouver, as one of Canada's largest cities, maintains a rich open dataset "
    "of crime incidents spanning over two decades. Predicting safety trends at the neighbourhood "
    "level can help municipal decision-makers allocate policing resources more effectively and "
    "enable residents to make informed choices about where to live, work, and commute. Despite "
    "the availability of this data, there is no publicly accessible tool that translates raw "
    "crime records into an interpretable, queryable safety score — this project fills that gap."
)

set_heading("1.2  Problem Definition", level=2)
p = doc.add_paragraph()
p.add_run("Input: ").font.bold = True
p.runs[-1].font.size = Pt(11)
p.add_run("Neighbourhood name and year (historical or future)").font.size = Pt(11)

p = doc.add_paragraph()
p.add_run("Output: ").font.bold = True
p.runs[-1].font.size = Pt(11)
p.add_run("Safety Score in the range [0, 100], where 100 is the safest").font.size = Pt(11)

p = doc.add_paragraph()
p.add_run("Task: ").font.bold = True
p.runs[-1].font.size = Pt(11)
p.add_run("Supervised regression — predicting a continuous safety score from crime features").font.size = Pt(11)

set_heading("1.3  Contributions", level=2)
bullet("Designed a weighted safety scoring system based on crime severity across 13 distinct crime types, with weights ranging from 1 (Mischief) to 10 (Homicide)")
bullet("Implemented a time-based train/test split to prevent data leakage, ensuring the model is evaluated on genuinely unseen future data")
bullet("Trained and compared three regression models (Linear Regression, Decision Tree, Random Forest) with full MSE, RMSE, and R² evaluation")
bullet("Built an iterative future-year prediction engine that chains predictions year-by-year beyond the dataset range")
bullet("Deployed the model through three interfaces: a CLI query tool, a terminal chatbot, and a GUI chatbot powered by a free local LLM (Ollama)")

doc.add_page_break()

# ── 2. Method ─────────────────────────────────────────────────────────────────
set_heading("2. Method")
set_heading("2.1  Algorithm Description", level=2)
body(
    "The safety score is derived from a weighted aggregation of crime incidents per "
    "neighbourhood-year. Let C = {c₁, c₂, ..., cₙ} be the set of crime records for a given "
    "neighbourhood n in year y, and let w(cᵢ) denote the severity weight of crime type cᵢ. "
    "The weighted crime score is:"
)
body("    weighted_score(n, y) = Σᵢ w(cᵢ)", bold=True)
body(
    "This is then min-max normalised and inverted across all neighbourhood-year pairs to "
    "produce a safety score in [0, 100]:"
)
body("    safety_score(n, y) = 100 × (1 − (weighted_score − min) / (max − min))", bold=True)
body(
    "where min and max are computed over the entire dataset. A score of 100 indicates the "
    "safest neighbourhood-year combination and 0 the most dangerous."
)
body(
    "Three regression models f(x; θ) are trained to predict safety_score from features x = "
    "[neighbourhood_enc, year, crime_count], minimising mean squared error:"
)
body("    L(θ) = (1/N) Σᵢ₌₁ᴺ (f(xᵢ; θ) − yᵢ)²", bold=True)
body(
    "where N is the number of training samples, xᵢ are the input features, and yᵢ is the "
    "true safety score."
)

set_heading("2.2  Overall Framework", level=2)
body("The pipeline proceeds through the following stages:")
steps = [
    "Raw Data Loading — Load VPD CSV (530,000+ records), retain only TYPE, YEAR, NEIGHBOURHOOD",
    "Data Cleaning — Drop rows with missing or blank NEIGHBOURHOOD values",
    "Crime Type Weighting — Map each crime type to a severity weight (1–10)",
    "Feature Engineering — Group by NEIGHBOURHOOD + YEAR, compute crime_count and weighted_score",
    "Safety Score Calculation — Min-max normalise weighted_score and invert to safety_score (0–100)",
    "Neighbourhood Encoding — Label-encode NEIGHBOURHOOD to a numeric ID",
    "Time-Based Split — Train: 2003–2019 (408 samples), Test: 2020–2023 (96 samples)",
    "Model Training — Fit Linear Regression, Decision Tree, Random Forest",
    "Evaluation — Compute MSE, RMSE, R² on the test set",
    "Model Saving — Save best model (Random Forest) to disk with joblib",
    "Prediction — Load saved model for CLI, chatbot, and GUI inference",
]
for i, step in enumerate(steps, 1):
    bullet(f"Step {i}: {step}")

set_heading("2.3  Design Choices", level=2)
body(
    "Features: We use neighbourhood identity (label-encoded), year, and crime count as features. "
    "The weighted_score is deliberately excluded from the feature set because safety_score is a "
    "direct mathematical transformation of it — including it would cause perfect linear fit and "
    "constitute data leakage. Year is included to capture temporal trends in crime patterns.",
    bold=False
)
doc.add_paragraph()
body(
    "Train/Test Split: A time-based split (train: 2003–2019, test: 2020–2023) is used instead "
    "of a random split. A random split would allow the model to see data from 2021 while "
    "predicting 2018, which is unrealistic and inflates performance metrics.",
    bold=False
)
doc.add_paragraph()
body(
    "Model Selection: Three models of increasing complexity are compared. Linear Regression "
    "serves as a baseline. Decision Tree captures nonlinear relationships but is prone to "
    "overfitting. Random Forest reduces variance through ensemble averaging, making it the "
    "most robust choice for this dataset.",
    bold=False
)
doc.add_paragraph()
body(
    "Future Prediction: For years beyond the dataset, predictions are chained iteratively. "
    "Each predicted year's score is used as input (prev_year_score, rolling_avg_score) for "
    "the next year, seeded from the last 3 known years of historical data.",
    bold=False
)

doc.add_page_break()

# ── 3. Experiment ─────────────────────────────────────────────────────────────
set_heading("3. Experiment")
set_heading("3.1  Experimental Setup", level=2)

p = doc.add_paragraph()
p.add_run("Dataset: ").font.bold = True
p.runs[-1].font.size = Pt(11)
p.add_run("Vancouver Police Department crime data sourced from Kaggle "
          "(https://www.kaggle.com/datasets/tcashion/vancouver-bc-crime-dataset). "
          "530,000+ individual crime incident records spanning 2003–2023 across "
          "24 Vancouver neighbourhoods.").font.size = Pt(11)

make_table(
    ["Statistic", "Value"],
    [
        ["Total records",          "530,000+"],
        ["Years covered",          "2003 – 2023"],
        ["Neighbourhoods",         "24"],
        ["Features used",          "TYPE, YEAR, NEIGHBOURHOOD"],
        ["Aggregated samples",     "504 (neighbourhood-year pairs)"],
        ["Training samples",       "408 (2003–2019)"],
        ["Test samples",           "96 (2020–2023)"],
        ["Target variable",        "Safety Score (0–100, continuous)"],
    ]
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Hyperparameters: ").font.bold = True
p.runs[-1].font.size = Pt(11)
p.add_run("Random Forest: n_estimators=200, random_state=42. "
          "Decision Tree: random_state=42. "
          "Linear Regression: default scikit-learn settings.").font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Software: ").font.bold = True
p.runs[-1].font.size = Pt(11)
p.add_run("Python 3.13, pandas 2.2.3, numpy 2.1.2, scikit-learn 1.8.0, "
          "matplotlib 3.10.7, joblib, ollama (llama3.2). CPU only.").font.size = Pt(11)

set_heading("3.2  Model Comparison & Results", level=2)
make_table(
    ["Model", "MSE", "RMSE", "R²"],
    [
        ["Linear Regression", "1.3802", "1.1748", "0.9923"],
        ["Decision Tree",     "1.0635", "1.0313", "0.9941"],
        ["Random Forest",     "1.0042", "1.0021", "0.9944"],
    ],
    best_row=2
)
doc.add_paragraph()
body(
    "Table 1: Performance comparison of all three models on the test set (2020–2023). "
    "Bold row indicates the best model. All models are evaluated using Mean Squared Error (MSE), "
    "Root Mean Squared Error (RMSE), and R² (coefficient of determination).",
    italic=True, size=9
)

set_heading("3.3  Justification & Analysis", level=2)
body(
    "Random Forest achieves the best performance across all three metrics (lowest MSE, lowest "
    "RMSE, highest R²). Its ensemble of 200 decision trees reduces variance through bootstrap "
    "aggregation, making it more robust to noise in the training data than a single Decision Tree."
)
body(
    "Decision Tree performs second, improving over Linear Regression by capturing nonlinear "
    "relationships between neighbourhood identity, year, and crime count. However, without "
    "ensemble averaging it is more sensitive to outliers in the training set."
)
body(
    "Linear Regression performs worst, as expected — the relationship between crime count and "
    "safety score is not purely linear across all neighbourhoods and years. Neighbourhoods with "
    "similar crime counts can have very different safety scores depending on the types of crimes "
    "committed, which a linear model cannot fully capture using only crime_count as a feature."
)
body(
    "All three models achieve R² > 0.99, indicating that the engineered features are highly "
    "predictive of the safety score. The primary source of error is in years with unusual crime "
    "spikes (e.g., 2020 COVID-19 period) where historical trends do not generalise well."
)
body(
    "Trade-offs: Random Forest is slower to train (200 trees) but is saved to disk after the "
    "first run, so subsequent loads are instantaneous. Linear Regression trains in milliseconds "
    "but sacrifices predictive accuracy. For this application, accuracy is prioritised over "
    "training speed since training is a one-time operation."
)

doc.add_page_break()

# ── 4. Conclusion ─────────────────────────────────────────────────────────────
set_heading("4. Conclusion")
set_heading("4.1  Summary", level=2)
body(
    "This project successfully built an end-to-end machine learning pipeline to predict safety "
    "scores for Vancouver neighbourhoods from historical VPD crime data. A weighted scoring "
    "system was designed to reflect crime severity, and three regression models were trained "
    "and evaluated using a time-based split to prevent data leakage. Random Forest achieved "
    "the best performance (RMSE = 1.0021, R² = 0.9944) and was saved as the production model. "
    "The model is deployed through three interfaces — a CLI tool, a terminal chatbot, and a "
    "GUI chatbot — all powered by a free local LLM (Ollama / llama3.2) that formats responses "
    "from real data without hallucinating facts."
)

set_heading("4.2  Limitations & Future Work", level=2)
body("Limitations:", bold=True)
bullet("The dataset may contain reporting bias — not all crimes are reported to the VPD, and reporting rates vary by neighbourhood and crime type.")
bullet("Only three features are used (neighbourhood, year, crime count). Socioeconomic factors such as population density, income levels, and housing costs are not considered.")
bullet("Label encoding of neighbourhoods assigns arbitrary numeric IDs and does not capture geographic proximity or similarity between adjacent neighbourhoods.")
bullet("Future year predictions are chained iteratively, meaning prediction error compounds over time — predictions for 2035 are less reliable than predictions for 2025.")
bullet("The model is trained on data up to 2023 and may not generalise well to structural changes in crime patterns beyond that point.")

doc.add_paragraph()
body("Future Work:", bold=True)
bullet("Incorporate demographic, population density, and socioeconomic data as additional features to improve prediction accuracy.")
bullet("Use geographic embeddings or spatial features (latitude/longitude centroids) instead of label encoding to capture neighbourhood proximity.")
bullet("Explore time-series models (LSTM, ARIMA) to better capture temporal trends and reduce compounding error in future year predictions.")
bullet("Apply hyperparameter tuning (grid search / random search) on Random Forest for further performance improvement.")
bullet("Extend the chatbot to support multi-turn reasoning, allowing users to ask follow-up questions that reference previous answers in the conversation.")

doc.add_page_break()

# ── 5. Distribution of Work ───────────────────────────────────────────────────
set_heading("5. Distribution of Work")
make_table(
    ["Team Member", "Responsibilities"],
    [
        ["Aakaash Senthilkumar (83091546)",
         "Model training, hyperparameter selection, future year prediction engine, "
         "chatbot and GUI development, pipeline integration, report writing"],
        ["Vaibhav Ambastha (19919539)",
         "Data preprocessing, feature engineering, crime type weighting system, "
         "safety score normalisation, data cleaning"],
        ["Arshvir Singh (62273818)",
         "Model evaluation, metric computation (MSE, RMSE, R²), result analysis, "
         "visualisation generation (7 figures), CLI query tool"],
    ]
)

doc.add_page_break()

# ── References ────────────────────────────────────────────────────────────────
set_heading("References")
refs = [
    "[1] T. Cashion. Vancouver BC Crime Dataset. Kaggle, 2024. "
    "https://www.kaggle.com/datasets/tcashion/vancouver-bc-crime-dataset",

    "[2] L. Breiman. Random Forests. Machine Learning, 45(1):5–32, 2001.",

    "[3] J. R. Quinlan. Induction of Decision Trees. Machine Learning, 1(1):81–106, 1986.",

    "[4] F. Pedregosa et al. Scikit-learn: Machine Learning in Python. "
    "Journal of Machine Learning Research, 12:2825–2830, 2011.",

    "[5] W. McKinney. Data Structures for Statistical Computing in Python. "
    "Proceedings of the 9th Python in Science Conference, 2010.",

    "[6] Ollama. Run Large Language Models Locally. https://ollama.com, 2024.",
]
for ref in refs:
    body(ref, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("report/CPEN355_Filled_Report.docx")
print("Report saved to report/CPEN355_Filled_Report.docx")
